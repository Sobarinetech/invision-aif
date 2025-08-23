import os
import streamlit as st
from google import genai
from google.genai import types
import tempfile
import PyPDF2
import docx
import pandas as pd
import plotly.express as px
import base64
from datetime import datetime
from supabase import create_client, Client

# ---------- CONFIG ----------
st.set_page_config(page_title="Invision AIF Solutions", layout="wide")
st.markdown(
    "<h1 style='color:#013a63;'>Invision AIF Solutions</h1>", unsafe_allow_html=True,
)
st.markdown(
    """
<div style='font-size:1.08rem;color:#222;background:#e9f5fe;border-radius:12px;padding:10px 18px;margin-bottom:1.5rem;'>
🔍 <b>Analyze, chat, and monitor with next-gen compliance AI for Alternative Investment Funds.<br> <span style='color:#013a63;'>Upload documents, ask questions, track insights with an advanced dashboard. Latest circulars and rules are always considered in the analysis.</span></b>
</div>
""",
    unsafe_allow_html=True,
)

# ---------- SUPABASE CLIENT ----------
# Accessing secrets directly from Render environment variables via st.secrets
SUPABASE_URL = st.secrets.get("SUPABASE_URL")
SUPABASE_KEY = st.secrets.get("SUPABASE_KEY")
CIRCULARS_COLUMNS = [
    "id",
    "guid",
    "title",
    "link",
    "description",
    "pub_date",
    "pdf_url",
    "type",
    "ai_analysis",
    "analysis_status",
    "created_at",
    "updated_at",
    "applicable_entities",
    "entity_keywords",
    "regulatory_scope",
    "extracted_content",
    "processing_status",
    "file_size",
    "extraction_metadata",
    "analyzed_at",
]

def get_supabase_client() -> Client:
    # Ensure SUPABASE_URL and SUPABASE_KEY are not None before creating client
    # This check will now specifically look for the values from st.secrets.get()
    if not SUPABASE_URL:
        st.error("Supabase URL (SUPABASE_URL) not found in Render environment variables.")
        st.stop() # Stop the app if credentials are missing
    if not SUPABASE_KEY:
        st.error("Supabase Key (SUPABASE_KEY) not found in Render environment variables.")
        st.stop() # Stop the app if credentials are missing
    return create_client(SUPABASE_URL, SUPABASE_KEY)

def fetch_latest_circulars(limit=5):
    try:
        client = get_supabase_client()
        data = (
            client.table("sebi_circulars")
            .select("*")
            .order("pub_date", desc=True)
            .limit(limit)
            .execute()
        )
        if hasattr(data, "data"):
            return data.data
    except Exception as e:
        st.error(f"Error fetching circulars from Supabase: {e}")
    return []

def make_circulars_context(circulars):
    if not circulars:
        return ""
    context = "Here are the most recent SEBI circulars and updates for regulatory context:\n"
    for c in circulars:
        context += (
            f"- {c.get('title','')}\n"
            f" Date: {c.get('pub_date','')}, Ref: {c.get('guid','')}\n"
            f" Description: {c.get('description','')}\n"
            f" [Full text]({c.get('link','')})\n"
        )
    context += "\n"
    return context

# ---------- GOOGLE GEMINI AI CONFIG ----------
def gemini_generate(input_text):
    gemini_api_key = st.secrets.get("GEMINI_API_KEY")
    if not gemini_api_key:
        st.error("Google Gemini API Key (GEMINI_API_KEY) not found in Render environment variables.")
        st.stop()

    client = genai.Client(
        api_key=gemini_api_key, # Correctly uses st.secrets.get()
    )
    model = "gemini-2.5-flash"
    contents = [
        types.Content(
            role="user",
            parts=[types.Part.from_text(text=input_text)],
        ),
    ]
    tools = [
        types.Tool(googleSearch=types.GoogleSearch()),
    ]
    generate_content_config = types.GenerateContentConfig(
        thinking_config=types.ThinkingConfig(
            thinking_budget=0,
        ),
        tools=tools,
    )
    output = ""
    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        if hasattr(chunk, "text") and chunk.text:
            output += chunk.text
        elif hasattr(chunk, "candidates") and chunk.candidates:
            part = chunk.candidates[0].content.parts[0]
            if part.text:
                output += part.text
    return output

def gemini_chat(history, doc_text=None, circulars_context=None):
    gemini_api_key = st.secrets.get("GEMINI_API_KEY")
    if not gemini_api_key:
        st.error("Google Gemini API Key (GEMINI_API_KEY) not found in Render environment variables.")
        st.stop()

    client = genai.Client(
        api_key=gemini_api_key, # Correctly uses st.secrets.get()
    )
    model = "gemini-2.5-flash"
    contents = []
    for idx, entry in enumerate(history):
        content = entry["content"]
        # Always prepend latest circulars context to the latest user message if idx == len(history) - 1 and entry["role"] == "user":
        if idx == len(history) - 1 and entry["role"] == "user":
            if circulars_context:
                content = f"{circulars_context}\n\n{content}"
            if doc_text:
                content += (
                    f"\n\n---\n\n"
                    f"Attached document text (for your analysis, reference, or to answer questions):\n"
                    f"{doc_text[:8000]}\n"
                )
        parts = [types.Part.from_text(text=content)]
        contents.append(types.Content(role="user" if entry["role"] == "user" else "model", parts=parts))
    tools = [
        types.Tool(googleSearch=types.GoogleSearch()),
    ]
    generate_content_config = types.GenerateContentConfig(
        thinking_config=types.ThinkingConfig(
            thinking_budget=0,
        ),
        tools=tools,
    )
    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        if hasattr(chunk, "text") and chunk.text:
            yield chunk.text
        elif hasattr(chunk, "candidates") and chunk.candidates:
            part = chunk.candidates[0].content.parts[0]
            if part.text:
                yield part.text

# ---------- FILE TEXT EXTRACTION ----------
def extract_text_from_pdf(pdf_file):
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(pdf_file.read())
        tmp_file.flush()
    file_path = tmp_file.name
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    finally:
        os.unlink(file_path)
    return text.strip()

def extract_text_from_docx(docx_file):
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(docx_file.read())
        tmp_file.flush()
    file_path = tmp_file.name
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    finally:
        os.unlink(file_path)
    return text.strip()

def extract_text_from_txt(txt_file):
    txt_file.seek(0)
    return txt_file.read().decode(errors="ignore").strip()

def extract_uploaded_files_text(uploaded_files):
    if not uploaded_files:
        return None
    all_texts = []
    for file in uploaded_files:
        ext = file.name.lower().split(".")[-1]
        if ext == "pdf":
            text = extract_text_from_pdf(file)
        elif ext == "docx":
            text = extract_text_from_docx(file)
        elif ext == "txt":
            text = extract_text_from_txt(file)
        else:
            text = ""
        if text:
            all_texts.append(f"---\nFile: {file.name}\n{text[:8000]}")
    if not all_texts:
        return None
    return "\n\n".join(all_texts)

def downloadable_report(report, filename="AIF-Compliance-Report.txt"):
    b64 = base64.b64encode(report.encode()).decode()
    return f'<a href="data:file/txt;base64,{b64}" download="{filename}" style="color:#00509e;">⬇️ Download Report</a>'

# ---------- LOCAL "DB" HELPER ----------
def get_dashboard_data():
    if "dashboard_data" not in st.session_state:
        st.session_state["dashboard_data"] = {
            "analyses": [],
            "chat_turns": 0,
            "chatbot_usage": [],
        }
    return st.session_state["dashboard_data"]

def save_dashboard_data(data):
    st.session_state["dashboard_data"] = data

# ---------- TABS ----------
tabs = st.tabs(
    [
        "Compliance Analysis",
        "RegOS Chatbot",
        "Dashboard & Insights",
        "SEBI Circulars Table",
    ]
)
dashboard_data = get_dashboard_data()

# 1. Compliance Analysis Tab
with tabs[0]:
    st.header("Compliance Analysis of AIF Documents")
    st.write(
        "Upload your AIF-related documents for direct compliance AI analysis. "
        "The AI will extract, analyze, and summarize compliance, risks, breaches, and more, referencing the latest regulations and providing citations. "
        "Recent SEBI circulars and regulatory updates are automatically blended into the analysis."
    )
    uploaded_files = st.file_uploader(
        "Upload document(s) (PDF, DOCX, TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="compliance_files",
    )

    # Fetch latest circulars
    latest_circulars = fetch_latest_circulars(limit=5)
    circulars_context = make_circulars_context(latest_circulars)

    if latest_circulars:
        with st.expander("Latest SEBI Circulars & Regulatory Updates used for analysis", expanded=False):
            for c in latest_circulars:
                st.markdown(f"**{c.get('title','')}** \nDate: {c.get('pub_date','')}, Ref: {c.get('guid','')}")
                st.markdown(f"Description: {c.get('description','')}")
                if c.get("link"):
                    st.markdown(f"[Full text]({c.get('link')})")
                st.markdown("---")

    if uploaded_files:
        doc_text = extract_uploaded_files_text(uploaded_files)
        if not doc_text:
            st.error(
                "No usable text extracted from your document(s). Please upload valid PDF, DOCX, or TXT files."
            )
        else:
            analysis_prompt = (
                f"{circulars_context}\n"
                "Act as a compliance expert. Analyze the following AIF (Alternative Investment Fund) document(s) strictly for compliance with current Indian regulations and regulatory context. "
                "For every point, provide proper grounded citations from official regulations, circulars, or law. Use real-time search and recent circulars above to ensure all referenced regulations are current. "
                "Return your analysis in the following structure:\n\n"
                "1. **Summary of Document(s)** (with citations)\n"
                "2. **Key Compliance Risks** (with grounded citations)\n"
                "3. **Detected Regulatory Breaches** (cite section/circular)\n"
                "4. **Recommendations** (reference the specific compliance to address)\n"
                "5. **Any Other Notable Observations** (with citations)\n"
                "6. **Breakdown of Risk by Section** (reference relevant requirements)\n"
                "7. **AI Confidence Level (0-100%) and Reasoning**\n"
                "8. **Potential Red Flags** (cite regulation)\n"
                "9. **Suggested Next Steps for Compliance Team** (with references)\n"
                "10. **All references must cite relevant regulation/circular with section number and date, if available.**\n\n"
                f"{doc_text}"
            )

            with st.spinner("AI analyzing uploaded document(s) with latest regulatory context..."):
                report = gemini_generate(analysis_prompt)
            st.subheader("Compliance Report (with citations)")
            st.markdown(report)
            st.markdown(
                downloadable_report(report, filename="AIF-Compliance-Report.txt"),
                unsafe_allow_html=True,
            )

            dashboard_data["analyses"].append(
                {
                    "name": ", ".join(f.name for f in uploaded_files),
                    "report": report,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                }
            )
            save_dashboard_data(dashboard_data)


# 2. RegOS Chatbot Tab (LLM with document upload)
with tabs[1]:
    st.header("RegOS Chatbot")
    st.write(
        "Ask regulatory, legal, or compliance questions about AIFs in India. "
        "AI is grounded in the latest regulatory context, blends in the most recent SEBI circulars, and always provides citations."
    )

    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []

    chat_uploaded_files = st.file_uploader(
        "Upload documents for this chat (optional, PDF, DOCX, TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="chat_files",
    )
    chat_doc_text = extract_uploaded_files_text(chat_uploaded_files) if chat_uploaded_files else None

    latest_circulars = fetch_latest_circulars(limit=5)
    circulars_context = make_circulars_context(latest_circulars)

    if latest_circulars:
        with st.expander("Latest SEBI Circulars & Regulatory Updates used for chatbot", expanded=False):
            for c in latest_circulars:
                st.markdown(f"**{c.get('title','')}** \nDate: {c.get('pub_date','')}, Ref: {c.get('guid','')}")
                st.markdown(f"Description: {c.get('description','')}")
                if c.get("link"):
                    st.markdown(f"[Full text]({c.get('link')})")
                st.markdown("---")

    # Show chat history as bubbles
    for entry in st.session_state["chat_history"]:
        if entry["role"] == "user":
            st.markdown(
                f"<div style='background:linear-gradient(90deg,#00509e,#2d82b7 60%); color:#fff; border-radius:16px; padding:12px 18px; margin-top:10px; margin-bottom:2px; max-width:85%; align-self:flex-end; margin-left:auto;'><b>You:</b> {entry['content']}</div>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f"<div style='background-color:#e9f5fe; color:#013a63; border-radius:16px; padding:12px 18px; margin-top:2px; margin-bottom:10px; max-width:85%; align-self:flex-start; margin-right:auto;'><b>RegOS AI:</b> {entry['content']}</div>",
                unsafe_allow_html=True,
            )

    with st.form(key="chat_form", clear_on_submit=True):
        user_input = st.text_input(
            "You:",
            key="chat_input",
            placeholder="Ask a regulatory, legal, or compliance question, or request analysis of your uploaded document...",
        )
        submitted = st.form_submit_button("Send")

        if submitted and user_input.strip():
            user_query = (
                f"{circulars_context}\n"
                "You are a compliance expert. Answer strictly with reference to the latest Indian regulations and regulatory context. "
                "Blend in recent SEBI circulars and official updates provided above. "
                "For every point, provide grounded citations from the current regulations, circulars, or law (include section number/date). "
                "Use real-time search to ensure all referenced regulations are current and provide links/citations where possible. "
                f"User query: {user_input}"
            )
            st.session_state["chat_history"].append({"role": "user", "content": user_query, "time": datetime.now().isoformat()})

            with st.spinner("RegOS AI (compliance expert) is typing..."):
                response_text = ""
                response_placeholder = st.empty()
                try:
                    for chunk in gemini_chat(st.session_state["chat_history"], doc_text=chat_doc_text, circulars_context=circulars_context):
                        response_text += chunk
                        response_placeholder.markdown(
                            f"<div style='background-color:#e9f5fe; color:#013a63; border-radius:16px; padding:12px 18px; margin-top:2px; margin-bottom:10px; max-width:85%; align-self:flex-start; margin-right:auto;'><b>RegOS AI:</b> {response_text}</div>",
                            unsafe_allow_html=True,
                        )
                except Exception as e:
                    response_text = f"Sorry, there was an error with the AI: {e}"
                    response_placeholder.markdown(response_text)

            st.session_state["chat_history"].append({"role": "model", "content": response_text, "time": datetime.now().isoformat()})

            dashboard_data["chat_turns"] += 1
            dashboard_data["chatbot_usage"].append(
                {
                    "prompt": user_input,
                    "response": response_text,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                }
            )
            save_dashboard_data(dashboard_data)
            st.rerun()

    if st.button("Clear Chat", key="clear_chat"):
        st.session_state["chat_history"] = []
        st.rerun()

    if st.button("Download Chat History"):
        chat_hist = "\n\n".join(
            f"{e['role'].capitalize()} ({e['time'][:19]}): {e['content']}"
            for e in st.session_state["chat_history"]
        )
        st.markdown(downloadable_report(chat_hist, filename="RegOS-Chat-History.txt"), unsafe_allow_html=True)


# 3. Dashboard & Insights Tab
with tabs[2]:
    st.header("Dashboard: Metrics, Trends & AI Insights")
    st.write(
        "Visualize compliance analysis results, AI performance metrics, chatbot usage, and key performance indicators. "
        "Track trends, download results, and get advanced breakdowns."
    )

    num_analyses = len(dashboard_data["analyses"])
    num_chats = dashboard_data["chat_turns"]

    st.metric("Documents Analyzed", num_analyses)
    st.metric("Chatbot Interactions", num_chats)

    uniq_files = set()
    for a in dashboard_data["analyses"]:
        for name in a["name"].split(","):
            uniq_files.add(name.strip())
    st.metric("Unique Files Uploaded", len(uniq_files))

    if dashboard_data["analyses"]:
        df_analyses = pd.DataFrame(dashboard_data["analyses"])
        df_analyses["timestamp"] = pd.to_datetime(df_analyses["timestamp"])
        fig = px.bar(
            df_analyses,
            x="timestamp",
            y=df_analyses.index + 1,
            hover_data=["name"],
            labels={"y": "Cumulative Analyses", "timestamp": "Time"},
            title="Document Analysis Timeline",
            color_discrete_sequence=["#00509e"],
        )
        st.plotly_chart(fig, use_container_width=True)

    if dashboard_data["chatbot_usage"]:
        df_chats = pd.DataFrame(dashboard_data["chatbot_usage"])
        df_chats["timestamp"] = pd.to_datetime(df_chats["timestamp"])
        chats_per_day = df_chats.groupby(df_chats["timestamp"].dt.date).size()
        fig2 = px.line(
            chats_per_day,
            markers=True,
            labels={"value": "# Interactions", "timestamp": "Date"},
            title="Chatbot Usage Over Time",
        )
        st.plotly_chart(fig2, use_container_width=True)

    st.subheader("Recent Analyses")
    for a in dashboard_data["analyses"][-3:][::-1]:
        with st.expander(f"{a['name']} ({a['timestamp']})"):
            st.write(a["report"])
            st.markdown(
                downloadable_report(a["report"], filename=f"{a['name']}-AIF-Compliance-Report.txt"),
                unsafe_allow_html=True,
            )

    st.subheader("Recent Chatbot Usage")
    for c in dashboard_data["chatbot_usage"][-3:][::-1]:
        with st.expander(f"Prompt: {c['prompt'][:50]}... ({c['timestamp']})"):
            st.markdown(f"**AI Response:** {c['response']}")

    st.info(
        "All data is stored in-memory for your session only. Download reports and chat logs for your records."
    )


# 4. SEBI Circulars Table Tab
with tabs[3]:
    st.header("SEBI Circulars Table")
    st.write("View the latest SEBI circulars and their details from the Supabase database.")
    # The get_supabase_client() function already handles potential missing credentials.
    # It will raise an error and stop the app if not found.
    client = get_supabase_client()
    circulars_data = (
        client.table("sebi_circulars").select("*").order("pub_date", desc=True).limit(50).execute()
    )
    if hasattr(circulars_data, "data") and circulars_data.data:
        df = pd.DataFrame(circulars_data.data)
        # Only show the columns requested
        display_cols = [c for c in CIRCULARS_COLUMNS if c in df.columns]
        st.dataframe(df[display_cols], use_container_width=True)
    else:
        st.info("No circulars found in the database.")

st.markdown("---")
st.caption("Powered by Google Gemini, Streamlit, Supabase, and Plotly. Confidential & Secure. 💡")

# ---------- REQUIREMENTS ----------
# pip install streamlit google-genai PyPDF2 python-docx pandas plotly supabase
