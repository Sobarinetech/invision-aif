import os
import streamlit as st
from google import genai
from google.genai import types
import tempfile
import PyPDF2
import docx
import pandas as pd
import plotly.express as px
import base64
from datetime import datetime
from supabase import create_client, Client

# Extra imports for Portfolio Company Monitoring
import requests
from googleapiclient.discovery import build
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from bs4 import BeautifulSoup, XMLParsedAsHTMLWarning
from langdetect import detect
import re
from textblob import TextBlob
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import warnings

# ---------- CONFIG ----------
st.set_page_config(page_title="Invision AIF Solutions", layout="wide")
st.markdown(
    "<h1 style='color:#013a63;'>Invision AIF Solutions</h1>",
    unsafe_allow_html=True,
)
st.markdown(
    """
    <div style='font-size:1.08rem;color:#222;background:#e9f5fe;border-radius:12px;padding:10px 18px;margin-bottom:1.5rem;'>
    🔍 <b>Analyze, chat, and monitor with next-gen compliance AI for Alternative Investment Funds.<br>
    <span style='color:#013a63;'>Upload documents, ask questions, track insights with an advanced dashboard. Latest circulars and rules are always considered in the analysis.</span></b>
    </div>
    """,
    unsafe_allow_html=True,
)

# ---------- SUPABASE CLIENT ----------
SUPABASE_URL = st.secrets.get("SUPABASE_URL")
SUPABASE_KEY = st.secrets.get("SUPABASE_KEY")

CIRCULARS_COLUMNS = [
    "id", "guid", "title", "link", "description", "pub_date", "pdf_url", "type", "ai_analysis", "analysis_status",
    "created_at", "updated_at", "applicable_entities", "entity_keywords", "regulatory_scope", "extracted_content",
    "processing_status", "file_size", "extraction_metadata", "analyzed_at"
]

def get_supabase_client() -> Client:
    return create_client(SUPABASE_URL, SUPABASE_KEY)

def fetch_latest_circulars(limit=5, table_name="sebi_circulars"):
    client = get_supabase_client()
    data = (
        client.table(table_name)
        .select("*")
        .order("pub_date", desc=True)
        .limit(limit)
        .execute()
    )
    if hasattr(data, "data"):
        return data.data
    return []

def make_circulars_context(circulars):
    if not circulars:
        return ""
    context = "Here are the most recent regulatory circulars and updates for context:\n"
    for c in circulars:
        context += (
            f"- {c.get('title','')}\n"
            f"  Date: {c.get('pub_date','')}, Ref: {c.get('guid','')}\n"
            f"  Description: {c.get('description','')}\n"
            f"  [Full text]({c.get('link','')})\n"
        )
    context += "\n"
    return context

# ---------- GOOGLE GEMINI AI CONFIG ----------
def gemini_generate(input_text):
    client = genai.Client(
        api_key=os.environ.get("GEMINI_API_KEY"),
    )
    model = "gemini-2.5-flash"
    contents = [
        types.Content(
            role="user",
            parts=[types.Part.from_text(text=input_text)],
        ),
    ]
    tools = [
        types.Tool(googleSearch=types.GoogleSearch()),
    ]
    generate_content_config = types.GenerateContentConfig(
        thinking_config=types.ThinkingConfig(
            thinking_budget=0,
        ),
        tools=tools,
    )
    output = ""
    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        if hasattr(chunk, "text") and chunk.text:
            output += chunk.text
        elif hasattr(chunk, "candidates") and chunk.candidates:
            part = chunk.candidates[0].content.parts[0]
            if part.text:
                output += part.text
    return output

def gemini_chat(history, doc_text=None, circulars_context=None):
    client = genai.Client(
        api_key=os.environ.get("GEMINI_API_KEY"),
    )
    model = "gemini-2.5-flash"
    contents = []
    for idx, entry in enumerate(history):
        content = entry["content"]
        # Always prepend latest circulars context to the latest user message
        if idx == len(history) - 1 and entry["role"] == "user":
            if circulars_context:
                content = f"{circulars_context}\n\n{content}"
            if doc_text:
                content += (
                    f"\n\n---\n\n"
                    f"Attached document text (for your analysis, reference, or to answer questions):\n"
                    f"{doc_text[:8000]}\n"
                )
        parts = [types.Part.from_text(text=content)]
        contents.append(types.Content(role="user" if entry["role"] == "user" else "model", parts=parts))
    tools = [
        types.Tool(googleSearch=types.GoogleSearch()),
    ]
    generate_content_config = types.GenerateContentConfig(
        thinking_config=types.ThinkingConfig(
            thinking_budget=0,
        ),
        tools=tools,
    )
    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        if hasattr(chunk, "text") and chunk.text:
            yield chunk.text
        elif hasattr(chunk, "candidates") and chunk.candidates:
            part = chunk.candidates[0].content.parts[0]
            if part.text:
                yield part.text

# New Edge Function for SEC RIA Compliance Analysis
def gemini_analyze_sec_compliance(doc_text, sec_circulars_context):
    client = genai.Client(
        api_key=os.environ.get("GEMINI_API_KEY"),
    )
    model = "gemini-2.5-flash"
    prompt = (
        f"{sec_circulars_context}\n"
        "Act as an SEC RIA compliance expert. Analyze the following document(s) strictly for compliance with current SEC RIA regulations. "
        "For every point, provide proper grounded citations from official SEC regulations, rules, or law. Use real-time search and recent circulars above to ensure all referenced regulations are current. "
        "Return your analysis in the following structure:\n\n"
        "1. **Summary of Document(s) in relation to SEC RIA Compliance** (with citations)\n"
        "2. **Key SEC RIA Compliance Risks** (with grounded citations)\n"
        "3. **Detected SEC RIA Regulatory Breaches** (cite rule/section)\n"
        "4. **Recommendations for SEC RIA Compliance** (reference the specific compliance to address)\n"
        "5. **Any Other Notable Observations for SEC RIA Compliance** (with citations)\n"
        "6. **Breakdown of Risk by Document Section against SEC RIA Rules** (reference relevant requirements)\n"
        "7. **AI Confidence Level (0-100%) and Reasoning**\n"
        "8. **Potential SEC RIA Red Flags** (cite regulation)\n"
        "9. **Suggested Next Steps for SEC RIA Compliance Team** (with references)\n"
        "10. **All references must cite relevant SEC regulation/rule with section number and date, if available.**\n\n"
        f"Document text for analysis:\n{doc_text}"
    )
    contents = [
        types.Content(
            role="user",
            parts=[types.Part.from_text(text=prompt)],
        ),
    ]
    tools = [
        types.Tool(googleSearch=types.GoogleSearch()),
    ]
    generate_content_config = types.GenerateContentConfig(
        thinking_config=types.ThinkingConfig(
            thinking_budget=0,
        ),
        tools=tools,
    )
    output = ""
    for chunk in client.models.generate_content_stream(
        model=model,
        contents=contents,
        config=generate_content_config,
    ):
        if hasattr(chunk, "text") and chunk.text:
            output += chunk.text
        elif hasattr(chunk, "candidates") and chunk.candidates:
            part = chunk.candidates[0].content.parts[0]
            if part.text:
                output += part.text
    return output


# ---------- FILE TEXT EXTRACTION ----------
def extract_text_from_pdf(pdf_file):
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(pdf_file.read())
        tmp_file.flush()
        file_path = tmp_file.name
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    finally:
        os.unlink(file_path)
    return text.strip()

def extract_text_from_docx(docx_file):
    text = ""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(docx_file.read())
        tmp_file.flush()
        file_path = tmp_file.name
    try:
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    finally:
        os.unlink(file_path)
    return text.strip()

def extract_text_from_txt(txt_file):
    txt_file.seek(0)
    return txt_file.read().decode(errors="ignore").strip()

def extract_uploaded_files_text(uploaded_files):
    if not uploaded_files:
        return None
    all_texts = []
    for file in uploaded_files:
        ext = file.name.lower().split(".")[-1]
        if ext == "pdf":
            text = extract_text_from_pdf(file)
        elif ext == "docx":
            text = extract_text_from_docx(file)
        elif ext == "txt":
            text = extract_text_from_txt(file)
        else:
            text = ""
        if text:
            all_texts.append(f"---\nFile: {file.name}\n{text[:8000]}")
    if not all_texts:
        return None
    return "\n\n".join(all_texts)

def downloadable_report(report, filename="AIF-Compliance-Report.txt"):
    b64 = base64.b64encode(report.encode()).decode()
    return f'<a href="data:file/txt;base64,{b64}" download="{filename}" style="color:#00509e;">⬇️ Download Report</a>'

# ---------- LOCAL "DB" HELPER ----------
def get_dashboard_data():
    if "dashboard_data" not in st.session_state:
        st.session_state["dashboard_data"] = {
            "analyses": [],
            "sec_ria_analyses": [], # Initialize sec_ria_analyses here
            "chat_turns": 0,
            "chatbot_usage": [],
        }
    # Ensure all keys exist, especially if session_state was partially initialized
    if "sec_ria_analyses" not in st.session_state["dashboard_data"]:
        st.session_state["dashboard_data"]["sec_ria_analyses"] = []
    if "analyses" not in st.session_state["dashboard_data"]:
        st.session_state["dashboard_data"]["analyses"] = []
    if "chat_turns" not in st.session_state["dashboard_data"]:
        st.session_state["dashboard_data"]["chat_turns"] = 0
    if "chatbot_usage" not in st.session_state["dashboard_data"]:
        st.session_state["dashboard_data"]["chatbot_usage"] = []

    return st.session_state["dashboard_data"]

def save_dashboard_data(data):
    st.session_state["dashboard_data"] = data

# ---------- TABS ----------
tabs = st.tabs(
    [
        "Compliance Analysis",
        "SEC RIA Compliance", # New tab
        "RegOS Chatbot",
        "Dashboard & Insights",
        "SEBI Circulars Table",
        "Portfolio Company Monitoring"
    ]
)
dashboard_data = get_dashboard_data()

# 1. Compliance Analysis Tab (Existing)
with tabs[0]:
    st.header("Compliance Analysis of AIF Documents")
    st.write(
        "Upload your AIF-related documents for direct compliance AI analysis. "
        "The AI will extract, analyze, and summarize compliance, risks, breaches, and more, referencing the latest regulations and providing citations. "
        "Recent SEBI circulars and regulatory updates are automatically blended into the analysis."
    )
    uploaded_files = st.file_uploader(
        "Upload document(s) (PDF, DOCX, TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="compliance_files"
    )
    # Fetch latest circulars
    latest_circulars = fetch_latest_circulars(limit=5, table_name="sebi_circulars")
    circulars_context = make_circulars_context(latest_circulars)
    if latest_circulars:
        with st.expander("Latest SEBI Circulars & Regulatory Updates used for analysis", expanded=False):
            for c in latest_circulars:
                st.markdown(f"**{c.get('title','')}**  \nDate: {c.get('pub_date','')}, Ref: {c.get('guid','')}")
                st.markdown(f"Description: {c.get('description','')}")
                if c.get("link"):
                    st.markdown(f"[Full text]({c.get('link')})")
                st.markdown("---")

    if uploaded_files:
        doc_text = extract_uploaded_files_text(uploaded_files)
        if not doc_text:
            st.error("No usable text extracted from your document(s). Please upload valid PDF, DOCX, or TXT files.")
        else:
            analysis_prompt = (
                f"{circulars_context}\n"
                "Act as a compliance expert. Analyze the following AIF (Alternative Investment Fund) document(s) strictly for compliance with current Indian regulations and regulatory context. "
                "For every point, provide proper grounded citations from official regulations, circulars, or law. Use real-time search and recent circulars above to ensure all referenced regulations are current. "
                "Return your analysis in the following structure:\n\n"
                "1. **Summary of Document(s)** (with citations)\n"
                "2. **Key Compliance Risks** (with grounded citations)\n"
                "3. **Detected Regulatory Breaches** (cite section/circular)\n"
                "4. **Recommendations** (reference the specific compliance to address)\n"
                "5. **Any Other Notable Observations** (with citations)\n"
                "6. **Breakdown of Risk by Section** (reference relevant requirements)\n"
                "7. **AI Confidence Level (0-100%) and Reasoning**\n"
                "8. **Potential Red Flags** (cite regulation)\n"
                "9. **Suggested Next Steps for Compliance Team** (with references)\n"
                "10. **All references must cite relevant regulation/circular with section number and date, if available.**\n\n"
                f"{doc_text}"
            )
            with st.spinner("AI analyzing uploaded document(s) with latest regulatory context..."):
                report = gemini_generate(analysis_prompt)
            st.subheader("Compliance Report (with citations)")
            st.markdown(report)
            st.markdown(downloadable_report(report, filename="AIF-Compliance-Report.txt"), unsafe_allow_html=True)
            dashboard_data["analyses"].append({
                "name": ", ".join(f.name for f in uploaded_files),
                "report": report,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            })
            save_dashboard_data(dashboard_data)

# 2. New SEC RIA Compliance Tab
with tabs[1]:
    st.header("SEC RIA Compliance Analysis")
    st.write(
        "Upload your documents for direct SEC Registered Investment Adviser (RIA) compliance AI analysis. "
        "The AI will analyze and summarize compliance, risks, and potential breaches against core SEC RIA obligations, "
        "referencing the latest SEC regulations and providing citations."
    )

    sec_uploaded_files = st.file_uploader(
        "Upload document(s) (PDF, DOCX, TXT) for SEC RIA compliance review",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="sec_ria_compliance_files"
    )

    # Simulate fetching SEC circulars/guidance from a theoretical "sec_circulars" table
    # In a real application, you would populate this table with actual SEC releases.
    # For now, it will return an empty list or mock data.
    sec_circulars = fetch_latest_circulars(limit=5, table_name="sec_circulars") # Assuming a table named sec_circulars
    sec_circulars_context = make_circulars_context(sec_circulars)
    if sec_circulars:
        with st.expander("Latest SEC Guidance & Regulatory Updates used for analysis", expanded=False):
            for c in sec_circulars:
                st.markdown(f"**{c.get('title','')}**  \nDate: {c.get('pub_date','')}, Ref: {c.get('guid','')}")
                st.markdown(f"Description: {c.get('description','')}")
                if c.get("link"):
                    st.markdown(f"[Full text]({c.get('link')})")
                st.markdown("---")
    else:
        st.info("No recent SEC circulars or guidance found in the database. Analysis will rely on general SEC RIA knowledge and real-time search.")


    st.subheader("Core SEC Compliance Obligations & Document Selection")
    st.markdown(
        """
        Below are key SEC RIA compliance obligations. Select the document type(s) your uploaded files pertain to.
        This helps the AI focus its analysis.
        """
    )

    sec_document_types = st.multiselect(
        "Select document type(s) for focused analysis:",
        [
            "Form ADV Part 1 & Part 2",
            "Form CRS (Client Relationship Summary)",
            "Compliance Program/Manual",
            "Code of Ethics",
            "Custody Rule related documents (e.g., audit reports)",
            "Marketing Materials/Advertising",
            "Books & Records (e.g., client agreements, trade records)",
            "Other Regulatory Filings",
            "General Internal Policies & Procedures"
        ],
        key="sec_doc_type_selector"
    )

    if sec_uploaded_files:
        sec_doc_text = extract_uploaded_files_text(sec_uploaded_files)
        if not sec_doc_text:
            st.error("No usable text extracted from your document(s). Please upload valid PDF, DOCX, or TXT files.")
        else:
            if st.button("Analyze SEC RIA Compliance"):
                with st.spinner("AI analyzing uploaded document(s) against SEC RIA regulations..."):
                    sec_ria_report = gemini_analyze_sec_compliance(sec_doc_text, sec_circulars_context)
                st.subheader("SEC RIA Compliance Report (with citations)")
                st.markdown(sec_ria_report)
                st.markdown(downloadable_report(sec_ria_report, filename="SEC-RIA-Compliance-Report.txt"), unsafe_allow_html=True)
                dashboard_data["sec_ria_analyses"].append({
                    "name": ", ".join(f.name for f in sec_uploaded_files),
                    "document_types": sec_document_types,
                    "report": sec_ria_report,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                })
                save_dashboard_data(dashboard_data)

# 3. RegOS Chatbot Tab (LLM with document upload)
with tabs[2]: # Changed index to 2
    st.header("RegOS Chatbot")
    st.write(
        "Ask regulatory, legal, or compliance questions about AIFs in India. "
        "AI is grounded in the latest regulatory context, blends in the most recent SEBI circulars, and always provides citations."
    )
    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []

    chat_uploaded_files = st.file_uploader(
        "Upload documents for this chat (optional, PDF, DOCX, TXT)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="chat_files"
    )
    chat_doc_text = extract_uploaded_files_text(chat_uploaded_files) if chat_uploaded_files else None

    latest_circulars = fetch_latest_circulars(limit=5, table_name="sebi_circulars")
    circulars_context = make_circulars_context(latest_circulars)
    if latest_circulars:
        with st.expander("Latest SEBI Circulars & Regulatory Updates used for chatbot", expanded=False):
            for c in latest_circulars:
                st.markdown(f"**{c.get('title','')}**  \nDate: {c.get('pub_date','')}, Ref: {c.get('guid','')}")
                st.markdown(f"Description: {c.get('description','')}")
                if c.get("link"):
                    st.markdown(f"[Full text]({c.get('link')})")
                st.markdown("---")

    # Show chat history as bubbles
    for entry in st.session_state["chat_history"]:
        if entry["role"] == "user":
            st.markdown(
                f"<div style='background:linear-gradient(90deg,#00509e,#2d82b7 60%); color:#fff; border-radius:16px; padding:12px 18px; margin-top:10px; margin-bottom:2px; max-width:85%; align-self:flex-end; margin-left:auto;'><b>You:</b> {entry['content']}</div>",
                unsafe_allow_html=True,
            )
        else:
            st.markdown(
                f"<div style='background-color:#e9f5fe; color:#013a63; border-radius:16px; padding:12px 18px; margin-top:2px; margin-bottom:10px; max-width:85%; align-self:flex-start; margin-right:auto;'><b>RegOS AI:</b> {entry['content']}</div>",
                unsafe_allow_html=True,
            )

    with st.form(key="chat_form", clear_on_submit=True):
        user_input = st.text_input(
            "You:",
            key="chat_input",
            placeholder="Ask a regulatory, legal, or compliance question, or request analysis of your uploaded document...",
        )
        submitted = st.form_submit_button("Send")

    if submitted and user_input.strip():
        user_query = (
            f"{circulars_context}\n"
            "You are a compliance expert. Answer strictly with reference to the latest Indian regulations and regulatory context. "
            "Blend in recent SEBI circulars and official updates provided above. "
            "For every point, provide grounded citations from the current regulations, circulars, or law (include section number/date). "
            "Use real-time search to ensure all referenced regulations are current and provide links/citations where possible. "
            f"User query: {user_input}"
        )
        st.session_state["chat_history"].append({"role": "user", "content": user_query, "time": datetime.now().isoformat()})
        with st.spinner("RegOS AI (compliance expert) is typing..."):
            response_text = ""
            response_placeholder = st.empty()
            try:
                for chunk in gemini_chat(st.session_state["chat_history"], doc_text=chat_doc_text, circulars_context=circulars_context):
                    response_text += chunk
                    response_placeholder.markdown(
                        f"<div style='background-color:#e9f5fe; color:#013a63; border-radius:16px; padding:12px 18px; margin-top:2px; margin-bottom:10px; max-width:85%; align-self:flex-start; margin-right:auto;'><b>RegOS AI:</b> {response_text}</div>",
                        unsafe_allow_html=True,
                    )
            except Exception as e:
                response_text = f"Sorry, there was an error with the AI: {e}"
                response_placeholder.markdown(response_text)
            st.session_state["chat_history"].append({"role": "model", "content": response_text, "time": datetime.now().isoformat()})
        dashboard_data["chat_turns"] += 1
        dashboard_data["chatbot_usage"].append({
            "prompt": user_input,
            "response": response_text,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        })
        save_dashboard_data(dashboard_data)
        st.rerun()

    if st.button("Clear Chat", key="clear_chat"):
        st.session_state["chat_history"] = []
        st.rerun()
    if st.button("Download Chat History"):
        chat_hist = "\n\n".join(
            f"{e['role'].capitalize()} ({e['time'][:19]}): {e['content']}" for e in st.session_state["chat_history"]
        )
        st.markdown(downloadable_report(chat_hist, filename="RegOS-Chat-History.txt"), unsafe_allow_html=True)

# 4. Dashboard & Insights Tab
with tabs[3]: # Changed index to 3
    st.header("Dashboard: Metrics, Trends & AI Insights")
    st.write(
        "Visualize compliance analysis results, AI performance metrics, chatbot usage, and key performance indicators. "
        "Track trends, download results, and get advanced breakdowns."
    )

    num_analyses = len(dashboard_data["analyses"])
    num_sec_ria_analyses = len(dashboard_data["sec_ria_analyses"]) # New metric
    num_chats = dashboard_data["chat_turns"]

    st.metric("AIF Documents Analyzed", num_analyses)
    st.metric("SEC RIA Documents Analyzed", num_sec_ria_analyses) # New metric display
    st.metric("Chatbot Interactions", num_chats)

    uniq_files = set()
    for a in dashboard_data["analyses"]:
        for name in a["name"].split(","):
            uniq_files.add(name.strip())
    for a in dashboard_data["sec_ria_analyses"]: # Include SEC RIA files in unique count
        for name in a["name"].split(","):
            uniq_files.add(name.strip())
    st.metric("Unique Files Uploaded (Total)", len(uniq_files))

    if dashboard_data["analyses"]:
        df_analyses = pd.DataFrame(dashboard_data["analyses"])
        df_analyses["timestamp"] = pd.to_datetime(df_analyses["timestamp"])
        fig = px.bar(
            df_analyses,
            x="timestamp",
            y=df_analyses.index+1,
            hover_data=["name"],
            labels={"y":"Cumulative AIF Analyses", "timestamp":"Time"},
            title="AIF Document Analysis Timeline",
            color_discrete_sequence=["#00509e"],
        )
        st.plotly_chart(fig, width='stretch') # Changed use_container_width to width='stretch'

    if dashboard_data["sec_ria_analyses"]:
        df_sec_ria_analyses = pd.DataFrame(dashboard_data["sec_ria_analyses"])
        df_sec_ria_analyses["timestamp"] = pd.to_datetime(df_sec_ria_analyses["timestamp"])
        fig_sec_ria = px.bar(
            df_sec_ria_analyses,
            x="timestamp",
            y=df_sec_ria_analyses.index+1,
            hover_data=["name", "document_types"],
            labels={"y":"Cumulative SEC RIA Analyses", "timestamp":"Time"},
            title="SEC RIA Document Analysis Timeline",
            color_discrete_sequence=["#28a745"], # Different color for distinction
        )
        st.plotly_chart(fig_sec_ria, width='stretch') # Changed use_container_width to width='stretch'


    if dashboard_data["chatbot_usage"]:
        df_chats = pd.DataFrame(dashboard_data["chatbot_usage"])
        df_chats["timestamp"] = pd.to_datetime(df_chats["timestamp"])
        chats_per_day = df_chats.groupby(df_chats["timestamp"].dt.date).size()
        fig2 = px.line(
            chats_per_day,
            markers=True,
            labels={"value":"# Interactions", "timestamp":"Date"},
            title="Chatbot Usage Over Time",
        )
        st.plotly_chart(fig2, width='stretch') # Changed use_container_width to width='stretch'

    st.subheader("Recent AIF Analyses")
    for a in dashboard_data["analyses"][-3:][::-1]:
        with st.expander(f"{a['name']} ({a['timestamp']})"):
            st.write(a["report"])
            st.markdown(downloadable_report(a["report"], filename=f"{a['name']}-AIF-Compliance-Report.txt"), unsafe_allow_html=True)

    st.subheader("Recent SEC RIA Analyses") # New section for SEC RIA
    for a in dashboard_data["sec_ria_analyses"][-3:][::-1]:
        with st.expander(f"{a['name']} (Types: {', '.join(a['document_types']) if a['document_types'] else 'N/A'}) ({a['timestamp']})"):
            st.write(a["report"])
            st.markdown(downloadable_report(a["report"], filename=f"{a['name']}-SEC-RIA-Compliance-Report.txt"), unsafe_allow_html=True)

    st.subheader("Recent Chatbot Usage")
    for c in dashboard_data["chatbot_usage"][-3:][::-1]:
        with st.expander(f"Prompt: {c['prompt'][:50]}... ({c['timestamp']})"):
            st.markdown(f"**AI Response:** {c['response']}")

    st.info("All data is stored in-memory for your session only. Download reports and chat logs for your records.")

# 5. SEBI Circulars Table Tab
with tabs[4]: # Changed index to 4
    st.header("SEBI Circulars Table")
    st.write("View the latest SEBI circulars and their details from the Supabase database.")
    client = get_supabase_client()
    circulars_data = client.table("sebi_circulars").select("*").order("pub_date", desc=True).limit(50).execute()
    if hasattr(circulars_data, "data") and circulars_data.data:
        df = pd.DataFrame(circulars_data.data)
        # Only show the columns requested
        display_cols = [c for c in CIRCULARS_COLUMNS if c in df.columns]
        st.dataframe(df[display_cols], width='stretch') # Changed use_container_width to width='stretch'
    else:
        st.info("No circulars found in the database.")

# 6. Portfolio Company Monitoring Tab
with tabs[5]: # Changed index to 5
    warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)
    st.header("Portfolio Company Monitoring")
    st.write(
        """
        Monitor news, mentions, and details for your Alternative Investment Fund portfolio companies.
        Search the entire web for news, events, or relevant updates. Enter company names below, and the tool will scan the web for recent mentions across news, blogs, and other sources.
        """
    )

    # Set up the Google API keys and Custom Search Engine ID
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    GOOGLE_SEARCH_ENGINE_ID = st.secrets["GOOGLE_SEARCH_ENGINE_ID"]

    # Initializing session state for detected matches
    if 'portfolio_matches' not in st.session_state:
        st.session_state.portfolio_matches = []

    # Input for list of portfolio companies
    company_list = st.text_area("Enter portfolio company names, one per line:", height=120, placeholder="Company A\nCompany B\nCompany C")
    company_names = [c.strip() for c in company_list.splitlines() if c.strip()]

    # Date filter (last N days)
    date_filter_days = st.slider("Restrict search to news from the past N days", 1, 60, 14)

    # Helper: Preprocess text
    def preprocess_text(text):
        text = text.lower()
        text = re.sub(r'[^a-zA-Z\s]', '', text)
        tokens = text.split()
        stop_words = set([
            'i', 'me', 'my', 'myself', 'we', 'our', 'ours', 'ourselves', 'you', 'your', 'yours', 'yourself', 'yourselves',
            'he', 'him', 'his', 'himself', 'she', 'her', 'hers', 'herself', 'it', 'its', 'itself', 'they', 'them', 'their',
            'theirs', 'themselves', 'what', 'which', 'who', 'whom', 'this', 'that', 'these', 'those', 'am', 'is', 'are',
            'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did', 'doing', 'a', 'an',
            'the', 'and', 'but', 'if', 'or', 'because', 'as', 'until', 'while', 'of', 'at', 'by', 'for', 'with', 'about',
            'against', 'between', 'into', 'through', 'during', 'before', 'after', 'above', 'below', 'to', 'from', 'up',
            'down', 'in', 'out', 'on', 'off', 'over', 'under', 'again', 'further', 'then', 'once', 'here', 'there', 'when',
            'where', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor',
            'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'can', 'will', 'just', 'don', 'should', 'now'
        ])
        filtered_tokens = [word for word in tokens if word not in stop_words]
        return " ".join(filtered_tokens)

    # Search and scan button
    if st.button("🔎 Scan News & Mentions for Portfolio Companies"):
        if not company_names:
            st.error("Please provide at least one company name.")
        else:
            with st.spinner('Scanning web for company news and mentions...'):
                service = build("customsearch", "v1", developerKey=GOOGLE_API_KEY)
                from datetime import timedelta
                import time as pytime
                import datetime as dt

                # Reset session state for detected matches
                st.session_state.portfolio_matches = []

                for company in company_names:
                    # Query for company, restrict to news and recent N days
                    search_query = f'"{company}"'
                    date_restrict = f"y[{(datetime.now() - timedelta(days=date_filter_days)).strftime('%Y%m%d')}]"  # dateRestrict=y[20230601]
                    try:
                        response = service.cse().list(
                            q=search_query,
                            cx=GOOGLE_SEARCH_ENGINE_ID,
                            num=8,
                            sort=f"date",
                            dateRestrict=date_restrict,
                            gl="in",  # India focus
                            cr="countryIN"
                        ).execute()
                        for result in response.get('items', []):
                            url = result['link']
                            title = result.get('title', '')
                            snippet = result.get('snippet', '')
                            pagedate = result.get('pagemap', {}).get('metatags', [{}])[0].get('article:published_time','')
                            try:
                                content_response = requests.get(url, timeout=10)
                                web_content = content_response.text if content_response.status_code == 200 else ""
                                soup = BeautifulSoup(web_content, "html.parser")
                                paragraphs = soup.find_all("p")
                                web_text = " ".join([para.get_text() for para in paragraphs])
                                # Sentiment
                                sentiment = ""
                                if web_text:
                                    tb = TextBlob(web_text[:2000])
                                    sentiment = tb.sentiment.polarity
                                # Detect Language
                                lang = ""
                                if web_text:
                                    try:
                                        lang = detect(web_text)
                                    except Exception:
                                        lang = ""
                                st.session_state.portfolio_matches.append({
                                    "Company": company,
                                    "URL": url,
                                    "Title": title,
                                    "Snippet": snippet[:240],
                                    "Date": pagedate[:10] if pagedate else "",
                                    "Language": lang,
                                    "Sentiment": sentiment,
                                })
                                pytime.sleep(1.0)  # avoid API throttling
                            except Exception:
                                continue
                    except Exception as e:
                        st.warning(f"Error searching for {company}: {e}")

            # Display dashboard results
            if st.session_state.portfolio_matches:
                st.success(f"News & mentions found for {len(st.session_state.portfolio_matches)} company-mentions!")
                df = pd.DataFrame(st.session_state.portfolio_matches)
                st.dataframe(df, width='stretch') # Changed use_container_width to width='stretch'

                # Show company-wise breakdown
                st.subheader("Company-wise Mentions Count")
                mentions_per_company = df['Company'].value_counts()
                fig = px.bar(
                    mentions_per_company,
                    labels={"value":"# Mentions", "Company":"Company"},
                    title="Mentions per Company",
                    color_discrete_sequence=["#5e35b1"]
                )
                st.plotly_chart(fig, width='stretch') # Changed use_container_width to width='stretch'

                # WordCloud of titles/snippets
                st.subheader("Word Cloud of News Headlines & Snippets")
                all_text = " ".join(df["Title"].fillna("").tolist() + df["Snippet"].fillna("").tolist())
                if all_text.strip():
                    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(all_text)
                    plt.figure(figsize=(8, 8))
                    plt.imshow(wordcloud, interpolation='bilinear')
                    plt.axis('off')
                    st.pyplot(plt)

                # Download option
                def convert_df(df):
                    return df.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📥 Download Mentions as CSV",
                    data=convert_df(df),
                    file_name="portfolio_mentions.csv",
                    mime="text/csv"
                )
            else:
                st.info("No news or web mentions found for your portfolio companies in the recent period.")

    st.caption("Uses Google Custom Search, BeautifulSoup, TextBlob, wordcloud, and Matplotlib. Results subject to Google CSE and web content restrictions.")

st.markdown("---")
st.caption("Powered by Google Gemini, Streamlit, Supabase, and Plotly. Confidential & Secure. 💡")
